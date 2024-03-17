import os
import tkinter as tk
from tkinter import filedialog

from docx import Document  # Document Conversion
from PyPDF2 import PdfFileReader # Document Conversion

def convert_doc():
    print("Welcome to the Document Conversion Tool!")
    
    # Browse for input file
    print("Please select the input document file to convert")
    root = tk.Tk()
    root.withdraw()
    input_path = filedialog.askopenfilename()
    root.destroy()  # Close the Tkinter window after file selection

    if not input_path:
        print("No input file selected. Please try again.")
        return

    output_format = input("Enter the output format (DOCX, PDF, TXT): ").upper()

    # Browse for the output folder
    print("Please select the destination folder to save the converted document file:")
    root = tk.Tk()
    root.withdraw()
    output_folder = filedialog.askdirectory()
    root.destroy()  # Close the Tkinter window after folder selection

    if not output_folder:
        print("No output folder selected. Please try again.")
        return
    
    supported_formats = ['DOCX', 'PDF', 'TXT']
    if output_format not in supported_formats:
        print("Unsupported output format.")
        return

    try:
        # Construct the output file path
        output_extension = output_format.lower()
        output_file = os.path.splitext(os.path.basename(input_path))[0] + '.' + output_extension
        output_path = os.path.join(output_folder, output_file)

        # Perform the document conversion
        if output_extension == 'docx':
            # Convert to DOCX (Assuming input is TXT file)
            with open(input_path, 'r') as txt_file:
                text_content = txt_file.read()
                doc = Document()
                doc.add_paragraph(text_content)
                doc.save(output_path)
        elif output_extension == 'pdf':
            # Convert to PDF (Assuming input is DOCX file)
            doc = Document(input_path)
            doc.save(output_path)
        elif output_extension == 'txt':
            # Convert to TXT (Assuming input is PDF file)
            with open(input_path, 'rb') as pdf_file:
                pdf_reader = PdfFileReader(pdf_file)
                text_content = ''
                for page_num in range(pdf_reader.numPages):
                    page = pdf_reader.getPage(page_num)
                    text_content += page.extractText()
                with open(output_path, 'w') as txt_output:
                    txt_output.write(text_content)

        print(f"Conversion successful: {os.path.basename(input_path)} -> {output_path}")
    except Exception as e:
        print(f"Conversion failed: {e}")
